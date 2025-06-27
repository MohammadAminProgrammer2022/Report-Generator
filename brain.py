from matplotlib import pyplot as plt
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt
import os
from random import randint
from docx.shared import Inches
from tqdm import tqdm


class DataReplacer:

    def __init__(self):
        self.counter = 0
        self.total_rows = 0
        self.days_file = False
        self.data_file = False

    def get_days(self, days_dir):
        '''
        input: day file directory
        output: 2 outputs one is boolean and the other is a message
        '''
        try:
            self.days = pd.read_excel(days_dir)
            self.days_file = True
            return True, 'فایل روزها با موفقیت دریافت شد'
        except Exception as e:
            print(e)
            msg = f'در دریافت فایل روزها خطایی وجود دارد. توجه داشته باشید حتما بایستی فایل شما در فالب اکسل باشد.:\n{e}'
            return False, msg

    def get_user_data(self, dir_data):
        '''
        input: data file directory
        output: 2 outputs one is boolean and the other is a message
        '''
        try:
            self.data = pd.read_excel(dir_data)
            self.data_file = True
            self.total_rows = len(self.data)
            return True, 'فایل اطلاعات حفاظ با موفقیت دریافت شد'
        except Exception as e:
            msg = f'در دریافت اطلاعات حفاظ خطایی وجود دارد. توجه داشته باشید حتما بایستی فایل شما در فالب اکسل باشد.:\n{e}'
            return False, msg

    @staticmethod
    def plotter(dict_data, valid_days):
        '''
        inputs: days and dictionary converted dataframe
        outputs: plot a graph and calculate sum and average of students' marks
        '''
        x = []
        y = []
        marks = []
        
        for _, m in dict_data.items():
            marks.append(m)
        marks = marks[4:]

        x_val = 1
        # print(valid_days)
        for m in marks:
            if len(valid_days) == len(x): break
            y.append(m)
            x.append(x_val)
            x_val += 1
        
        # average
        y_sum = sum(y)
        avg = y_sum / len(y)
        avg = round(avg, 2)

        plt.figure(figsize=(12, 6))
        plt.plot(x, y, marker='o', linestyle='-', color='blue', markerfacecolor='red', markeredgecolor='red')

        # Set x and y axis ticks explicitly
        plt.xticks(ticks=range(1, 30))       # Show all 1 to 31
        plt.yticks(ticks=range(0, 101, 10))  # Show every 10 from 0 to 100

        fig_name = 'fig.png'

        plt.grid(True)
        plt.tight_layout()
        plt.savefig(fig_name)
        # plt.show()
        
        # image path
        img_path = os.path.join(os.getcwd(), fig_name)
        return img_path, y_sum, avg

    @staticmethod
    def accumulator(dict_data, dict_days, image, total, average):
        """
        This function is responssible to create
        a dictionary of all the data of a student.
        image: an image path of the graph
        """
        data = {}
        data.update(dict_data)
        data.update(dict_days)
        data['image'] = image
        data['total'] = total # sum of all marks
        data['average'] = average

        return data

    @staticmethod
    def raplace_data(total_data, src='data.docx'):
        folder_name = f'گزارشات{"_" + str(total_data['month']) + "_" + str(total_data['year'])}'
        try:
            os.mkdir(folder_name)
        except:
            pass
        save_path = os.path.join(os.getcwd(), folder_name)
        doc = Document(src)

        for table in doc.tables:
            # print(table)
            for row in table.rows:
                for cell in row.cells:
                    # Iterate over paragraphs in the cell
                    for para in cell.paragraphs:
                        for run in para.runs:
                            # print(f'run: {run.text}')
                            if run.text == 'image':
                                run.text = run.text.replace('image', '')
                                run = para.add_run()
                                run.add_picture(total_data['image'], width=Inches(10.37), height=Inches(4.69))
                                # h: 3.12 , w: 10.73

                            for key, value in total_data.items():
                                if key == run.text:
                                    run.text = run.text.replace(key, str(value))

        # Save the modified document
        file_name = f"{total_data['year']}_{total_data['month']}_{total_data['name']}_{randint(0, 999999)}.docx"
        # output_path = Path(folder_name) / file_name
        doc.save(os.path.join(save_path, file_name))

    @staticmethod
    def df2dict(df):
        """
        This function receives a dataframe and returns
        a dictionary in a list of that dataframe.
        """
        return df.to_dict('records')

    @staticmethod
    def days_validator(days):
        '''
        This function drops null values
        and returns valid data
        '''
        valid_days = days.dropna(axis=1)
        valid_days = DataReplacer.df2dict(valid_days)[0]
        return valid_days

    def data_replace_main(self, template_path):
        self.counter += 1
        data = self.data
        days = self.days
        # print('in data rep')

        dict_data = DataReplacer.df2dict(data)
        # print(dict_data)
        valid_days = DataReplacer.days_validator(days)
        # print(valid_days)
        # # valid_days

        for data in tqdm(dict_data):
            yield True
            # print('s')
            img, sum_score, average_score = DataReplacer.plotter(data, valid_days)
            total_data = DataReplacer.accumulator(data, valid_days, img, sum_score, average_score)
            DataReplacer.raplace_data(total_data, src=template_path)
        
        self.data_file, self.days_file = False, False
        try: 
            os.remove('fig.png')
        except:
            pass
